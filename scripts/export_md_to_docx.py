#!/usr/bin/env python3
"""将项目内 Markdown 章节导出为 Word（.docx），供论文排版使用。

优先使用 python-docx（版式更完整）；未安装时自动用标准库生成符合 OOXML 的 .docx。
"""
from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

try:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt
    from docx.oxml.ns import qn

    _HAS_DOCX = True
except ImportError:
    Document = None  # type: ignore
    WD_LINE_SPACING = None  # type: ignore
    Pt = None  # type: ignore
    qn = None  # type: ignore
    _HAS_DOCX = False

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _xml_escape_text(s: str) -> str:
    return escape(s, {"'": "&apos;", '"': "&quot;"})


def _ooxml_run(text: str, *, bold: bool = False, sz_half_pt: int = 24, font_east_asia: str = "宋体") -> str:
    rpr = f'<w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="{font_east_asia}"/><w:sz w:val="{sz_half_pt}"/><w:szCs w:val="{sz_half_pt}"/>'
    if bold:
        rpr += "<w:b/><w:bCs/>"
    rpr += "</w:rPr>"
    # Word 要求 w:t 中含空格/换行时加 xml:space="preserve"
    t = _xml_escape_text(text)
    preserve = ' xml:space="preserve"' if (text.startswith(" ") or text.endswith(" ") or "  " in text) else ""
    return f"<w:r>{rpr}<w:t{preserve}>{t}</w:t></w:r>"


def _ooxml_paragraph_from_parts(runs_xml: str) -> str:
    return (
        "<w:p>"
        '<w:pPr><w:spacing w:line="360" w:lineRule="auto"/></w:pPr>'
        f"{runs_xml}</w:p>"
    )


def _ooxml_heading(text: str, level: int) -> str:
    sizes = {1: 32, 2: 28, 3: 24}
    sz = sizes.get(level, 24)
    runs = _ooxml_run(text, bold=True, sz_half_pt=sz, font_east_asia="黑体")
    return _ooxml_paragraph_from_parts(runs)


def _ooxml_body_from_mixed_line(line: str) -> str:
    parts = re.split(r"(\*\*[^*]+\*\*)", line)
    chunks: list[str] = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            chunks.append(_ooxml_run(part[2:-2], bold=True))
        else:
            chunks.append(_ooxml_run(part, bold=False))
    return _ooxml_paragraph_from_parts("".join(chunks))


def md_to_docx_ooxml(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    body_parts: list[str] = []
    for line in lines:
        line = line.rstrip()
        if not line or line.strip() == "---":
            continue
        if line.startswith("# "):
            body_parts.append(_ooxml_heading(line[2:].strip(), 1))
        elif line.startswith("## "):
            body_parts.append(_ooxml_heading(line[3:].strip(), 2))
        elif line.startswith("### "):
            body_parts.append(_ooxml_heading(line[4:].strip(), 3))
        else:
            body_parts.append(_ooxml_body_from_mixed_line(line))

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:document xmlns:w="{_W_NS}"><w:body>'
        + "".join(body_parts)
        + "<w:sectPr><w:pgSz w:w=\"11906\" w:h=\"16838\"/><w:pgMar w:top=\"1440\" w:right=\"1440\" "
        + "w:bottom=\"1440\" w:left=\"1440\"/></w:sectPr></w:body></w:document>"
    )

    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    doc_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>
"""

    out_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/document.xml", document_xml)
        z.writestr("word/_rels/document.xml.rels", doc_rels)


def _set_run_font(run, western: str = "Times New Roman", east_asia: str = "宋体", size_pt: float = 12):
    run.font.name = western
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east_asia)


def add_mixed_paragraph(doc, text: str, *, bold_base: bool = False, italic: bool = False, size_pt: float = 12):
    """将 **粗体** 解析为 Word 粗体 runs。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
            if bold_base:
                pass
            _set_run_font(run, size_pt=size_pt)
        else:
            run = p.add_run(part)
            run.bold = bold_base
            run.italic = italic
            _set_run_font(run, size_pt=size_pt)
    return p


def md_to_docx(md_path: Path, out_path: Path, *, title: str | None = None) -> None:
    if not _HAS_DOCX:
        if title:
            print("警告: 无 python-docx 时忽略 --title，改用 OOXML 简易导出", file=sys.stderr)
        md_to_docx_ooxml(md_path, out_path)
        print(f"已写入（标准库 OOXML）: {out_path}")
        return

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        _set_run_font(r, east_asia="黑体", size_pt=18)
        doc.add_paragraph()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        if not line:
            continue
        if line.strip() == "---":
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for r in h.runs:
                _set_run_font(r, east_asia="黑体", size_pt=16)
            continue
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for r in h.runs:
                _set_run_font(r, east_asia="黑体", size_pt=14)
            continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            for r in h.runs:
                _set_run_font(r, east_asia="黑体", size_pt=12)
            continue

        add_mixed_paragraph(doc, line, size_pt=12)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"已写入: {out_path}")


def main():
    ap = argparse.ArgumentParser(description="Markdown → Word docx")
    ap.add_argument("input", type=Path, help="输入 .md 路径")
    ap.add_argument("-o", "--output", type=Path, help="输出 .docx 路径（默认同名）")
    ap.add_argument("--title", default=None, help="可选封面式标题（level 0）")
    args = ap.parse_args()
    inp = args.input.resolve()
    out = args.output
    if out is None:
        out = inp.with_suffix(".docx")
    else:
        out = out.resolve()
    md_to_docx(inp, out, title=args.title)


if __name__ == "__main__":
    main()
